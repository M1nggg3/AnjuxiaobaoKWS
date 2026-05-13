from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs"
ASSET_DIR = OUT_DIR / "assets"
DOCX_PATH = OUT_DIR / "安居小宝语音唤醒技术实现路线.docx"
FLOWCHART_PATH = ASSET_DIR / "anjuxiaobao_kws_flowchart.png"


TITLE_COLOR = RGBColor(31, 78, 121)
ACCENT_COLOR = "1F4E79"
LIGHT_BLUE = "D9EAF7"
LIGHT_GREEN = "E2F0D9"
LIGHT_GRAY = "F2F2F2"


def font_path():
    candidates = [
        Path(r"C:\Windows\Fonts\msyh.ttc"),
        Path(r"C:\Windows\Fonts\simsun.ttc"),
        Path(r"C:\Windows\Fonts\simhei.ttf"),
    ]
    for path in candidates:
        if path.exists():
            return str(path)
    return None


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, size=10, color=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if bold else WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def style_table(table, header_fill=ACCENT_COLOR):
    table.style = "Table Grid"
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Microsoft YaHei"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                    run.font.size = Pt(9.5)
            if i == 0:
                set_cell_shading(cell, header_fill)
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
            else:
                set_cell_shading(cell, "FFFFFF" if i % 2 else "F8FBFD")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_widths(table, widths_cm):
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            if idx < len(row.cells):
                row.cells[idx].width = Cm(width)


def add_heading(document, text, level):
    paragraph = document.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.color.rgb = TITLE_COLOR
    return paragraph


def add_body_paragraph(document, text, bold_prefix=None):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.18
    if bold_prefix and text.startswith(bold_prefix):
        run = paragraph.add_run(bold_prefix)
        run.bold = True
        rest = text[len(bold_prefix):]
        paragraph.add_run(rest)
    else:
        paragraph.add_run(text)
    for run in paragraph.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(10.5)
    return paragraph


def add_bullet(document, text):
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10.2)


def create_flowchart():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    img = Image.new("RGB", (1800, 980), "white")
    draw = ImageDraw.Draw(img)
    fp = font_path()
    title_font = ImageFont.truetype(fp, 52) if fp else ImageFont.load_default()
    box_font = ImageFont.truetype(fp, 31) if fp else ImageFont.load_default()
    small_font = ImageFont.truetype(fp, 25) if fp else ImageFont.load_default()

    draw.text((70, 45), "“安居小宝”语音唤醒项目流程图", fill=(31, 78, 121), font=title_font)

    boxes = [
        ((90, 170, 460, 310), "需求定义", "关键词、指标、板端约束"),
        ((560, 170, 930, 310), "数据采集", "正样本、负样本、噪声场景"),
        ((1030, 170, 1400, 310), "数据整理", "wav.scp / text / data.list"),
        ((560, 400, 930, 540), "PC 训练验证", "WeKWS FSMN-CTC 微调"),
        ((1030, 400, 1400, 540), "评估调参", "召回率、误唤醒、阈值"),
        ((90, 640, 460, 780), "模型导出", "TorchScript / ONNX / 端侧格式"),
        ((560, 640, 930, 780), "RK3566 集成", "录音链路、推理、唤醒回调"),
        ((1030, 640, 1400, 780), "实机验收", "远近场、噪声、长期稳定性"),
    ]

    def rounded_box(rect, title, subtitle, fill, outline):
        draw.rounded_rectangle(rect, radius=22, fill=fill, outline=outline, width=4)
        x1, y1, x2, y2 = rect
        tw = draw.textbbox((0, 0), title, font=box_font)
        sw = draw.textbbox((0, 0), subtitle, font=small_font)
        draw.text((x1 + (x2 - x1 - (tw[2] - tw[0])) / 2, y1 + 32), title, fill=(35, 55, 75), font=box_font)
        draw.text((x1 + (x2 - x1 - (sw[2] - sw[0])) / 2, y1 + 86), subtitle, fill=(80, 88, 96), font=small_font)

    for idx, (rect, title, subtitle) in enumerate(boxes):
        fill = (217, 234, 247) if idx in (0, 3, 4) else (226, 240, 217) if idx in (5, 6, 7) else (242, 242, 242)
        outline = (31, 78, 121) if idx in (0, 3, 4) else (84, 130, 53) if idx in (5, 6, 7) else (120, 120, 120)
        rounded_box(rect, title, subtitle, fill, outline)

    arrows = [
        ((460, 240), (560, 240)),
        ((930, 240), (1030, 240)),
        ((1215, 310), (1215, 400)),
        ((1030, 470), (930, 470)),
        ((745, 540), (745, 640)),
        ((560, 710), (460, 710)),
        ((460, 710), (560, 710)),
        ((930, 710), (1030, 710)),
    ]
    for start, end in arrows:
        draw.line([start, end], fill=(90, 90, 90), width=5)
        ex, ey = end
        sx, sy = start
        if ex > sx:
            points = [(ex, ey), (ex - 18, ey - 12), (ex - 18, ey + 12)]
        elif ex < sx:
            points = [(ex, ey), (ex + 18, ey - 12), (ex + 18, ey + 12)]
        elif ey > sy:
            points = [(ex, ey), (ex - 12, ey - 18), (ex + 12, ey - 18)]
        else:
            points = [(ex, ey), (ex - 12, ey + 18), (ex + 12, ey + 18)]
        draw.polygon(points, fill=(90, 90, 90))

    draw.text((90, 860), "闭环原则：实机采集问题样本 → 回流训练集 → 重训/调阈值 → 再验收", fill=(120, 80, 20), font=small_font)
    img.save(FLOWCHART_PATH, quality=95)


def setup_document():
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    styles = document.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10.5)
    return document


def build_docx():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    create_flowchart()
    document = setup_document()

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    run = title.add_run("“安居小宝”语音唤醒功能\n技术实现路线")
    run.bold = True
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(24)
    run.font.color.rgb = TITLE_COLOR

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("基于 WeKWS / FSMN-CTC，先完成 PC 训练验证，再推进 RK3566 端侧移植")
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(90, 90, 90)

    meta = document.add_table(rows=3, cols=2)
    meta.style = "Table Grid"
    meta.autofit = False
    set_table_widths(meta, [3.2, 12.6])
    rows = [
        ("项目目标", "实现“安居小宝”本地语音唤醒，形成可训练、可评估、可移植的工程闭环。"),
        ("阶段范围", "本阶段聚焦需求拆解、数据准备、PC 端训练验证、端侧移植方案设计。"),
        ("目标平台", "RK3566 开发板，后续结合实际麦克风链路完成端侧推理和实机验证。"),
    ]
    for row, values in zip(meta.rows, rows):
        set_cell_text(row.cells[0], values[0], bold=True, size=9.5, color="FFFFFF")
        set_cell_text(row.cells[1], values[1], size=9.5)
        set_cell_shading(row.cells[0], ACCENT_COLOR)
        set_cell_shading(row.cells[1], "FFFFFF")

    add_heading(document, "一、总体技术路线", 1)
    add_body_paragraph(document, "项目采用“PC 端训练验证 + RK3566 端侧移植”的两阶段路线。PC 阶段先完成数据集规范、模型训练、评估和阈值选择；端侧阶段再围绕录音链路、特征提取、模型推理和唤醒回调完成集成。")
    add_body_paragraph(document, "模型路线：参考 WeKWS 的 FSMN-CTC 关键词检测方案，使用“安居小宝”作为目标唤醒词，优先复用已有预训练模型进行微调，以降低首轮训练成本和数据压力。")

    add_heading(document, "二、关键模块划分", 1)
    module_table = document.add_table(rows=1, cols=4)
    module_table.autofit = False
    set_table_widths(module_table, [3.5, 4.1, 4.0, 4.2])
    headers = ["模块", "主要工作", "阶段产物", "风险关注"]
    for idx, header in enumerate(headers):
        module_table.rows[0].cells[idx].text = header
    module_rows = [
        ("数据采集与整理", "采集正负样本，统一 16k 单通道 WAV，生成 wav.scp、text、data.list", "训练集、验证集、测试集", "人员覆盖不足、板端麦克风数据缺失"),
        ("PC 训练验证", "基于 WeKWS 训练 FSMN-CTC 模型，完成冒烟测试和正式训练", "模型权重、训练日志、评估报告", "正负样本比例、阈值稳定性"),
        ("模型评估调参", "统计召回率、误唤醒率、不同噪声场景表现，选择部署阈值", "阈值建议、问题样本清单", "测试集和真实场景不一致"),
        ("RK3566 端侧集成", "接入录音、特征、推理、唤醒事件，完成实机测试", "端侧 demo、部署说明", "算力、延迟、音频链路差异"),
    ]
    for values in module_rows:
        row = module_table.add_row()
        for idx, value in enumerate(values):
            row.cells[idx].text = value
    set_table_widths(module_table, [3.5, 4.1, 4.0, 4.2])
    style_table(module_table)

    document.add_page_break()
    add_heading(document, "三、项目流程图", 1)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(FLOWCHART_PATH), width=Cm(15.8))
    caption = document.add_paragraph("图 1  “安居小宝”语音唤醒项目流程")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in caption.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(90, 90, 90)

    add_heading(document, "四、阶段实施计划", 1)
    phase_table = document.add_table(rows=1, cols=5)
    phase_table.autofit = False
    set_table_widths(phase_table, [3.0, 2.2, 3.1, 3.1, 4.4])
    for idx, header in enumerate(["阶段", "周期建议", "输入", "输出", "通过标准"]):
        phase_table.rows[0].cells[idx].text = header
    phase_rows = [
        ("1. 冒烟验证", "1 周", "少量正负样本、参考配置", "可运行训练脚本和初版模型", "训练能收敛，模型可完成关键词检测"),
        ("2. 小规模验证", "2-3 周", "2k-5k 正样本、20h+ 负样本", "评估报告和阈值范围", "召回率、误唤醒趋势满足继续投入条件"),
        ("3. 正式训练", "3-6 周", "2.5w+ 正样本、100h+ 负样本", "正式模型和测试报告", "独立测试集指标达到项目目标"),
        ("4. 端侧移植", "2-4 周", "正式模型、板端录音链路", "RK3566 demo 和部署包", "延迟、内存、稳定性通过实机验收"),
    ]
    for values in phase_rows:
        row = phase_table.add_row()
        for idx, value in enumerate(values):
            row.cells[idx].text = value
    set_table_widths(phase_table, [3.0, 2.2, 3.1, 3.1, 4.4])
    style_table(phase_table)

    add_heading(document, "五、资源与协作需求", 1)
    add_bullet(document, "数据资源：协调录音人员、环境噪声、远近场距离、最终设备麦克风链路，保证训练数据贴近真实使用场景。")
    add_bullet(document, "算力资源：PC 训练阶段建议准备可用 GPU 环境，冒烟测试可低配，正式训练建议独占或稳定排期。")
    add_bullet(document, "测试资源：需要建立独立测试集，并在 RK3566 实机上长期采集误唤醒和漏唤醒问题样本。")
    add_bullet(document, "工程资源：需要音频采集、模型训练、端侧集成、产品测试共同配合，形成数据回流闭环。")

    add_heading(document, "六、当前建议", 1)
    add_body_paragraph(document, "建议先申请一轮冒烟测试资源：完成 ADB 取数、少量样本整理、WeKWS 训练脚本跑通和 RK3566 录音链路验证。该阶段成本低，但能尽早暴露数据格式、录音质量、模型导出和端侧集成风险。")

    document.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build_docx()
